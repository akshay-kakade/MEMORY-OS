from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import pandas as pd
import io
import os

class ExportService:
    def export_as_pdf(self, chat_title, messages):
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        width, height = letter
        
        p.setFont("Helvetica-Bold", 16)
        p.drawString(100, height - 50, f"Chat Export: {chat_title}")
        
        p.setFont("Helvetica", 12)
        y = height - 80
        for msg in messages:
            role = msg['role'].capitalize()
            content = msg['content']
            
            # Simple text wrapping for PDF
            lines = [f"{role}: {content[i:i+80]}" for i in range(0, len(content), 80)]
            for line in lines:
                if y < 50:
                    p.showPage()
                    y = height - 50
                p.drawString(100, y, line)
                y -= 15
            y -= 10
            
        p.save()
        buffer.seek(0)
        return buffer.getvalue()

    def export_as_docx(self, chat_title, messages):
        doc = Document()
        doc.add_heading(f"Chat Export: {chat_title}", 0)
        
        for msg in messages:
            p = doc.add_paragraph()
            p.add_run(f"{msg['role'].capitalize()}: ").bold = True
            p.add_run(msg['content'])
            
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def export_as_excel(self, messages):
        df = pd.DataFrame(messages)
        buffer = io.BytesIO()
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='ChatHistory')
        buffer.seek(0)
        return buffer.getvalue()

    def export_as_csv(self, messages):
        df = pd.DataFrame(messages)
        return df.to_csv(index=False).encode('utf-8')

export_service = ExportService()
